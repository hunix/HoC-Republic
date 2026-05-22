#!/usr/bin/env python3
"""
HoC Professional Word Document Generator v2
Templates: research-paper, policy, report, memo, proposal
Features: cover page, TOC, headers/footers, styled tables, images, branding
"""
import json, os, sys, datetime, urllib.request, subprocess

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Parse args from JSON file ───────────────────────────────────
_args_path = "/tmp/_gen_args.json"
if os.path.exists(_args_path):
    with open(_args_path, "r") as f:
        args = json.load(f)
elif len(sys.argv) > 1:
    args = json.loads(sys.argv[1])
else:
    args = {}
slide_data_raw = args.get("slide_data", "[]")
branding_raw = args.get("branding", "{}")
title_text = args.get("title", "Untitled Document")
out_path = args.get("out_path", "/workspace/document.docx")
template_name = args.get("template", "report")
images_raw = args.get("images", "[]")

try:    sections = json.loads(slide_data_raw)
except: sections = []
if isinstance(sections, dict):
    sections = sections.get("sections", sections.get("slides", []))

try:    brand_data = json.loads(branding_raw)
except: brand_data = {}

try:    images_list = json.loads(images_raw)
except: images_list = []

# ── SVG→PNG Conversion ─────────────────────────────────────────
def svg_to_png(svg_path, png_path=None, width=800):
    if not png_path: png_path = svg_path.rsplit(".", 1)[0] + ".png"
    try:
        import cairosvg
        cairosvg.svg2png(url=svg_path, write_to=png_path, output_width=width)
        return png_path
    except ImportError: pass
    try:
        r = subprocess.run(["rsvg-convert", "-w", str(width), svg_path, "-o", png_path],
                           capture_output=True, timeout=10)
        if r.returncode == 0: return png_path
    except: pass
    return None

def download_image(url, dest):
    try:
        urllib.request.urlretrieve(url, dest)
        is_svg = dest.lower().endswith(".svg")
        if not is_svg:
            with open(dest, "rb") as f:
                is_svg = b"<svg" in f.read(256)
        if is_svg:
            png = svg_to_png(dest, dest.rsplit(".", 1)[0] + ".png")
            if png: return png
        return dest
    except: return None

# ── Template Configs ────────────────────────────────────────────
TEMPLATES = {
    "research-paper": {
        "accent": "#1a365d", "accent_rgb": (26, 54, 93),
        "heading_font": "Times New Roman", "body_font": "Times New Roman",
        "body_size": 12, "heading_sizes": [24, 18, 14, 12],
        "has_cover": True, "has_toc": True, "has_abstract": True,
        "page_margins": (1.0, 1.0, 1.0, 1.0),
    },
    "policy": {
        "accent": "#1e40af", "accent_rgb": (30, 64, 175),
        "heading_font": "Calibri", "body_font": "Calibri",
        "body_size": 11, "heading_sizes": [22, 16, 13, 11],
        "has_cover": True, "has_toc": True, "has_abstract": False,
        "page_margins": (1.0, 1.0, 1.0, 1.0),
    },
    "report": {
        "accent": "#0f172a", "accent_rgb": (15, 23, 42),
        "heading_font": "Calibri", "body_font": "Calibri",
        "body_size": 11, "heading_sizes": [26, 18, 14, 12],
        "has_cover": True, "has_toc": True, "has_abstract": False,
        "page_margins": (1.0, 1.0, 1.0, 1.0),
    },
    "memo": {
        "accent": "#374151", "accent_rgb": (55, 65, 81),
        "heading_font": "Arial", "body_font": "Arial",
        "body_size": 11, "heading_sizes": [20, 16, 13, 11],
        "has_cover": False, "has_toc": False, "has_abstract": False,
        "page_margins": (1.0, 1.0, 1.0, 1.0),
    },
    "proposal": {
        "accent": "#7c3aed", "accent_rgb": (124, 58, 237),
        "heading_font": "Calibri", "body_font": "Calibri",
        "body_size": 11, "heading_sizes": [28, 18, 14, 12],
        "has_cover": True, "has_toc": True, "has_abstract": False,
        "page_margins": (1.0, 1.0, 1.0, 1.0),
    },
}

tpl = TEMPLATES.get(template_name, TEMPLATES["report"])
ACCENT = tpl["accent"]
ACCENT_RGB = tpl["accent_rgb"]
COMPANY = brand_data.get("company_name", "")
LOGO_URL = brand_data.get("logo_url", "")
TODAY = datetime.date.today().strftime("%B %d, %Y")

# ── Document Setup ──────────────────────────────────────────────
doc = Document()

# Set default font
style = doc.styles["Normal"]
font = style.font
font.name = tpl["body_font"]
font.size = Pt(tpl["body_size"])
font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Set page margins
for section in doc.sections:
    m = tpl["page_margins"]
    section.top_margin = Inches(m[0])
    section.right_margin = Inches(m[1])
    section.bottom_margin = Inches(m[2])
    section.left_margin = Inches(m[3])

# ── Style Heading Formats ───────────────────────────────────────
for level in range(1, 5):
    style_name = f"Heading {level}"
    if style_name in doc.styles:
        hs = doc.styles[style_name]
        hs.font.name = tpl["heading_font"]
        hs.font.size = Pt(tpl["heading_sizes"][level-1])
        hs.font.color.rgb = RGBColor(*ACCENT_RGB)
        hs.font.bold = True
        hs.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
        hs.paragraph_format.space_after = Pt(8)

# ── Helper Functions ────────────────────────────────────────────
def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:color="{ACCENT.lstrip("#")}"/></w:pBdr>')
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(12)

def add_styled_table(doc, headers, rows, accent_rgb=None):
    """Add a professionally styled table."""
    if not accent_rgb: accent_rgb = ACCENT_RGB
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row styling
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = str(h)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{ACCENT.lstrip("#")}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(10)

    # Data rows with alternating colors
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            if ri % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="f8fafc"/>')
                cell._tc.get_or_add_tcPr().append(shading)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()  # spacing

def add_rich_content(doc, content):
    """Parse content with bullets, numbered lists, and paragraphs."""
    lines = str(content).split("\\n")
    for line in lines:
        stripped = line.strip()
        if not stripped: continue
        # Numbered list
        if stripped[:2] in ("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9."):
            p = doc.add_paragraph(stripped[2:].strip(), style="List Number")
        # Bullet point
        elif stripped.startswith(("-", "•", "*", "▸")):
            text = stripped.lstrip("-•*▸ ")
            p = doc.add_paragraph(text, style="List Bullet")
        # Blockquote
        elif stripped.startswith(">"):
            p = doc.add_paragraph(stripped[1:].strip())
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].font.italic = True if p.runs else None
            p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b) if p.runs else None
        else:
            p = doc.add_paragraph(stripped)

# ── Cover Page ──────────────────────────────────────────────────
if tpl["has_cover"]:
    # Logo
    if LOGO_URL:
        logo = download_image(LOGO_URL, "/tmp/_docx_logo.png")
        if logo and os.path.exists(logo):
            try: doc.add_picture(logo, width=Inches(1.5))
            except: pass

    # Spacer
    for _ in range(3):
        doc.add_paragraph()

    # Accent bar (simulated with paragraph border)
    add_horizontal_rule(doc)

    # Title
    title_p = doc.add_heading(title_text, 0)
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title_p.runs:
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(*ACCENT_RGB)

    # Subtitle
    sub_text = ""
    if sections and sections[0].get("subtitle"):
        sub_text = sections[0]["subtitle"]
    elif template_name == "policy":
        sub_text = "Policy & Procedures Document"
    elif template_name == "research-paper":
        sub_text = "Research Paper"
    elif template_name == "proposal":
        sub_text = "Project Proposal"
    if sub_text:
        sp = doc.add_paragraph(sub_text)
        sp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sp.runs[0].font.size = Pt(16)
        sp.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    add_horizontal_rule(doc)

    # Metadata
    meta_items = []
    if COMPANY: meta_items.append(f"Organization: {COMPANY}")
    meta_items.append(f"Date: {TODAY}")
    meta_items.append(f"Template: {template_name.replace('-', ' ').title()}")

    for item in meta_items:
        mp = doc.add_paragraph(item)
        mp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in mp.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    doc.add_page_break()

# ── Table of Contents Placeholder ───────────────────────────────
if tpl["has_toc"]:
    doc.add_heading("Table of Contents", level=1)
    toc_p = doc.add_paragraph("(Update field to generate — right-click → Update Field in Word)")
    toc_p.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
    toc_p.runs[0].font.italic = True

    # Add TOC field code (Word auto-generates on open)
    fld_p = doc.add_paragraph()
    run = fld_p.add_run()
    fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar)
    run2 = fld_p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._r.append(instrText)
    run3 = fld_p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)

    doc.add_page_break()

# ── Abstract (research-paper only) ──────────────────────────────
if tpl["has_abstract"] and sections:
    abstract = ""
    for s in sections:
        if s.get("layout") == "abstract" or "abstract" in s.get("title", "").lower():
            abstract = s.get("content", "")
            break
    if abstract:
        doc.add_heading("Abstract", level=1)
        ap = doc.add_paragraph(abstract)
        ap.paragraph_format.left_indent = Inches(0.5)
        ap.paragraph_format.right_indent = Inches(0.5)
        for run in ap.runs:
            run.font.italic = True
        add_horizontal_rule(doc)

# ── Memo Header ─────────────────────────────────────────────────
if template_name == "memo":
    doc.add_heading("MEMORANDUM", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_horizontal_rule(doc)
    memo_fields = {"TO": "", "FROM": COMPANY or "", "DATE": TODAY, "RE": title_text}
    for s in sections:
        if s.get("layout") == "memo_header":
            memo_fields.update({"TO": s.get("to", ""), "FROM": s.get("from", COMPANY),
                                "CC": s.get("cc", ""), "RE": s.get("subject", title_text)})
            break
    for k, v in memo_fields.items():
        if v:
            p = doc.add_paragraph()
            run_label = p.add_run(f"{k}:  ")
            run_label.bold = True
            p.add_run(v)
    add_horizontal_rule(doc)

# ── Build Sections ──────────────────────────────────────────────
for si, section in enumerate(sections):
    layout = section.get("layout", "content")
    s_title = section.get("title", "")
    s_content = section.get("content", "")

    # Skip abstract/memo_header (already handled)
    if layout in ("abstract", "memo_header"): continue

    # Section heading
    if s_title and layout != "title":
        level = section.get("level", 1)
        level = min(max(level, 1), 4)
        doc.add_heading(s_title, level=level)

    # Content
    if layout == "content" or layout == "section":
        add_rich_content(doc, s_content)

    elif layout == "table":
        headers = section.get("headers", [])
        rows = section.get("rows", [])
        if not headers and s_content:
            for line in s_content.split("\\n"):
                cells = [c.strip() for c in line.split("|") if c.strip()]
                if cells:
                    if not headers: headers = cells
                    else: rows.append(cells)
        if headers:
            add_styled_table(doc, headers, rows)
        else:
            add_rich_content(doc, s_content)

    elif layout == "image":
        img_url = section.get("image_url", "")
        caption = section.get("caption", "")
        if img_url:
            img_path = download_image(img_url, f"/tmp/_docx_img_{si}.png")
            if img_path and os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(5))
                    last_p = doc.paragraphs[-1]
                    last_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                except: pass
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cp.runs[0].font.italic = True
            cp.runs[0].font.size = Pt(9)
            cp.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
        if s_content:
            add_rich_content(doc, s_content)

    elif layout == "two_column":
        parts = s_content.split("|||") if "|||" in str(s_content) else [str(s_content), ""]
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, part in enumerate(parts[:2]):
            cell = table.rows[0].cells[ci]
            p = cell.paragraphs[0]
            p.text = part.strip()
            for run in p.runs:
                run.font.size = Pt(tpl["body_size"])

    elif layout in ("stats", "kpi"):
        stats = section.get("stats", [])
        if not stats and s_content:
            for part in s_content.split("|"):
                kv = part.strip().split(":")
                if len(kv) >= 2:
                    stats.append({"label": kv[0].strip(), "value": kv[1].strip()})
        if stats:
            ncols = min(len(stats), 4)
            table = doc.add_table(rows=2, cols=ncols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for si_s, stat in enumerate(stats[:ncols]):
                # Value row
                cell = table.rows[0].cells[si_s]
                p = cell.paragraphs[0]
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(str(stat.get("value", "—")))
                run.font.size = Pt(24); run.bold = True
                run.font.color.rgb = RGBColor(*ACCENT_RGB)
                # Label row
                cell2 = table.rows[1].cells[si_s]
                p2 = cell2.paragraphs[0]
                p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run2 = p2.add_run(stat.get("label", ""))
                run2.font.size = Pt(10)
                run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
            doc.add_paragraph()

    elif layout == "callout":
        # Indented callout box
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        icon = section.get("icon", "ℹ️")
        run = p.add_run(f"  {icon}  {s_content}")
        run.font.italic = True
        run.font.color.rgb = RGBColor(*ACCENT_RGB)

    elif layout == "references":
        doc.add_heading("References", level=1)
        refs = section.get("references", [])
        if not refs and s_content:
            refs = [r.strip() for r in s_content.split("\\n") if r.strip()]
        for ri, ref in enumerate(refs):
            p = doc.add_paragraph(f"[{ri+1}]  {ref}")
            p.paragraph_format.hanging_indent = Inches(0.5)
            p.runs[0].font.size = Pt(10)

    elif layout == "appendix":
        doc.add_page_break()
        doc.add_heading(f"Appendix: {s_title}", level=1)
        add_rich_content(doc, s_content)

    else:
        add_rich_content(doc, s_content)

# ── Add Images ──────────────────────────────────────────────────
for img in images_list:
    url = img.get("url", "")
    caption = img.get("caption", "")
    if url:
        path = download_image(url, f"/tmp/_docx_extra_{images_list.index(img)}.png")
        if path and os.path.exists(path):
            try:
                doc.add_picture(path, width=Inches(5))
                doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except: pass
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cp.runs[0].font.italic = True
            cp.runs[0].font.size = Pt(9)

# ── Headers & Footers ──────────────────────────────────────────
for section in doc.sections:
    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = hp.add_run(f"{COMPANY}  |  {title_text}" if COMPANY else title_text)
    run.font.size = Pt(8); run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fp.add_run(f"{TODAY}  •  Confidential")
    run.font.size = Pt(8); run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

# ── Save ────────────────────────────────────────────────────────
doc.save(out_path)
print(f"Document saved: {out_path} (template: {template_name}, {len(sections)} sections)")
